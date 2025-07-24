import os
import logging
from typing import Dict, Any
import re
from datetime import datetime
from pathlib import Path

# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import markdown
from io import BytesIO

logger = logging.getLogger(__name__)

def generate_docx(report_content: str, product_id: str) -> str:
    """
    Generate a Word document from the PSUR report content
    
    Args:
        report_content: Markdown formatted report content
        product_id: Product ID for file naming
    
    Returns:
        Path to the generated DOCX file
    """
    
    try:
        # Create output directory if it doesn't exist
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        
        # Create filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"PSUR_Report_{product_id}_{timestamp}.docx"
        file_path = output_dir / filename
        
        # Create new document
        document = Document()
        
        # Set document properties
        document.core_properties.title = f"PSUR Report - Product {product_id}"
        document.core_properties.author = "Pharma Pulse System"
        document.core_properties.subject = "Periodic Safety Update Report"
        
        # Add custom styles
        add_custom_styles(document)
        
        # Parse markdown content and add to document
        parse_markdown_to_docx(document, report_content)
        
        # Save document
        document.save(str(file_path))
        
        logger.info(f"DOCX report generated: {file_path}")
        return str(file_path)
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise Exception(f"Failed to generate Word document: {str(e)}")

def add_custom_styles(document: Document):
    """Add custom styles to the document"""
    
    styles = document.styles
    
    # Main heading style
    if 'CustomTitle' not in [style.name for style in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(18)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Section heading style
    if 'CustomHeading' not in [style.name for style in styles]:
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(14)
        heading_font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
    
    # Normal text style
    if 'CustomNormal' not in [style.name for style in styles]:
        normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)

def parse_markdown_to_docx(document: Document, markdown_content: str):
    """Parse markdown content and add formatted content to Word document"""
    
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Add empty paragraph for spacing
            document.add_paragraph()
            continue
        
        # Main title (single #)
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            p = document.add_paragraph(title_text, style='CustomTitle')
            
        # Section headings (##)
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            p = document.add_paragraph(heading_text, style='CustomHeading')
            
        # Sub-headings (###)
        elif line.startswith('### '):
            subheading_text = line[4:].strip()
            p = document.add_paragraph(subheading_text)
            p.style.font.bold = True
            p.style.font.size = Pt(12)
            
        # Table rows
        elif line.startswith('|') and '|' in line[1:]:
            # Handle table - this is simplified, real implementation would need table parsing
            table_text = line.replace('|', ' | ').strip()
            p = document.add_paragraph(table_text, style='CustomNormal')
            
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            p = document.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            
        # Horizontal rule
        elif line.startswith('---'):
            document.add_paragraph('_' * 50, style='CustomNormal')
            
        # Regular text
        else:
            # Remove markdown formatting for simple implementation
            clean_text = clean_markdown_text(line)
            if clean_text:
                document.add_paragraph(clean_text, style='CustomNormal')

def clean_markdown_text(text: str) -> str:
    """Clean markdown formatting from text"""
    
    # Remove bold formatting
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Remove italic formatting
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Remove code formatting
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    return text.strip()

def generate_pdf(report_content: str, product_id: str) -> str:
    """
    Generate a PDF document from the PSUR report content
    
    Args:
        report_content: Markdown formatted report content
        product_id: Product ID for file naming
    
    Returns:
        Path to the generated PDF file
    """
    
    try:
        # Create output directory if it doesn't exist
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        
        # Create filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"PSUR_Report_{product_id}_{timestamp}.pdf"
        file_path = output_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            fontName='Helvetica'
        )
        
        # Parse content and create flowables
        story = []
        story.extend(parse_markdown_to_pdf(report_content, title_style, heading_style, normal_style))
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"PDF report generated: {file_path}")
        return str(file_path)
        
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise Exception(f"Failed to generate PDF document: {str(e)}")

def parse_markdown_to_pdf(content: str, title_style, heading_style, normal_style) -> list:
    """Parse markdown content and return list of PDF flowables"""
    
    story = []
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            story.append(Spacer(1, 6))
            continue
        
        # Main title (single #)
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 12))
            
        # Section headings (##)
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            story.append(Paragraph(heading_text, heading_style))
            
        # Sub-headings (###)
        elif line.startswith('### '):
            subheading_text = line[4:].strip()
            sub_style = ParagraphStyle(
                'SubHeading',
                parent=normal_style,
                fontSize=12,
                fontName='Helvetica-Bold',
                spaceBefore=8,
                spaceAfter=4
            )
            story.append(Paragraph(subheading_text, sub_style))
            
        # Tables (simplified)
        elif line.startswith('|') and '|' in line[1:]:
            # Convert table row to paragraph for simplicity
            table_text = line.replace('|', ' | ').strip()
            story.append(Paragraph(table_text, normal_style))
            
        # Horizontal rule
        elif line.startswith('---'):
            story.append(Spacer(1, 6))
            separator_style = ParagraphStyle(
                'Separator',
                parent=normal_style,
                alignment=1
            )
            story.append(Paragraph('_' * 50, separator_style))
            story.append(Spacer(1, 6))
            
        # Regular text
        else:
            clean_text = clean_markdown_text(line)
            if clean_text:
                # Handle bold text in PDF
                if '**' in clean_text:
                    clean_text = clean_text.replace('**', '<b>').replace('**', '</b>')
                story.append(Paragraph(clean_text, normal_style))
    
    return story

def create_table_from_markdown(table_lines: list) -> Table:
    """Create a ReportLab table from markdown table lines"""
    
    data = []
    
    for line in table_lines:
        if line.startswith('|') and line.endswith('|'):
            # Remove leading and trailing pipes, split by pipes
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            data.append(cells)
    
    if not data:
        return None
    
    # Create table
    table = Table(data)
    
    # Apply table style
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    return table

def get_export_statistics() -> Dict[str, int]:
    """Get statistics about exported files"""
    
    output_dir = Path("output")
    
    if not output_dir.exists():
        return {"docx_files": 0, "pdf_files": 0, "total_files": 0}
    
    docx_files = len(list(output_dir.glob("*.docx")))
    pdf_files = len(list(output_dir.glob("*.pdf")))
    total_files = len(list(output_dir.glob("*")))
    
    return {
        "docx_files": docx_files,
        "pdf_files": pdf_files,
        "total_files": total_files
    }

def cleanup_old_files(days_old: int = 7):
    """Clean up old exported files"""
    
    try:
        output_dir = Path("output")
        if not output_dir.exists():
            return
        
        current_time = datetime.now()
        
        for file_path in output_dir.glob("*"):
            if file_path.is_file():
                file_age = current_time - datetime.fromtimestamp(file_path.stat().st_mtime)
                if file_age.days > days_old:
                    file_path.unlink()
                    logger.info(f"Deleted old file: {file_path}")
        
    except Exception as e:
        logger.error(f"Error cleaning up old files: {str(e)}")
